import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("Missing dependency: python-docx. Install with: pip install python-docx")


CONTENT = {
    "title": "Grookai Vault — Investor Brief",
    "sections": [
        ("Executive Summary", (
            "Grookai Vault is wired end-to-end across Supabase (database schema + Edge functions), Flutter UI (ThunderShell app), and DevOps scripts.",
            "Core flows exist (Search, Vault tracking, public Wall groundwork, Scanner dev tools). The system is ~70% production-ready.",
            "The fastest path to a stable beta is aligning a few naming/schema mismatches between migrations, Edge functions, and Flutter queries, then standardizing types and run/dev tooling.",
        )),
        ("What’s Completed", (
            "Backend schema foundations: public.vault_items with RLS; RPC public.vault_add_item(...); staged listings/listing_images with RLS; materialized view + view for feed; refresh RPC; unified search view with trigram/unaccent indexing.",
            "Edge functions: wall_feed (needs view name alignment), search_cards, importers/cron engines, health probes.",
            "Flutter application: App shell/login/tabbed nav; Wall grid/infinite scroll; Search, Card Detail, Vault list, Profile; env via flutter_dotenv with fallback.",
            "DevOps/scripts: repair/pull/push and diagnostics helpers; documented device run checklist.",
        )),
        ("What’s Next (High-Impact Fixes)", (
            "Unify Wall feed naming to public.wall_feed_view across DB/Edge/Flutter.",
            "Apply Wall base tables from _hold; validate RLS + refresh RPC sequencing.",
            "Fix Search image field via image_best alias (or app selects).",
            "Align Vault schema vs app (qty, condition_label, grade_label) or reduce insert payload.",
            "Consolidate on listings; archive legacy wall_posts functions.",
            "Ensure grants; add typegen (TS) and DTOs/codegen (Flutter).",
            "Surface errors with SnackBars; maintain layout hygiene.",
        )),
        ("Current Readiness (Self-Assessment)", (
            "Database Schema: 70% | RLS/Grants: 75% | Edge Functions: 65% | Flutter UI: 75% | DevOps/Tooling: 70% | Overall E2E: 70%",
        )),
        ("Comparison To A High-End Dev Team", (
            "Architecture on par; naming/contract hygiene slightly below due to drift.",
            "Type safety below top-tier (no generated types for Edge; limited DTOs in Flutter).",
            "Dev ergonomics solid; add all-in-one run/stop tasks.",
            "Test coverage not observed; premium teams add RLS/integration tests.",
            "Net: ~70% of premium output today; quick wins can lift to 85–90%.",
        )),
        ("Cost To Reach Production-Ready Beta", (
            "Lean contractors ($80–$120/hr): 2–3 engineers x 2–3 weeks → $25k–$55k.",
            "High-end shop ($150–$220/hr): 2–3 engineers + PM x 2–3 weeks → $60k–$120k.",
            "Assumes no major surprises with RLS/data; hosted keys/secrets available; on-device testing ready.",
        )),
        ("30/60/90 Plan", (
            "Next 7–10 days: unify wall_feed_view; promote base tables; add image_best; align vault_items or app.",
            "Days 11–30: consolidate on listings; add TS typegen + Flutter DTO/codegen; VS Code tasks; SnackBar errors.",
            "Days 31–90: add integration tests; monitoring for importers/pricing; seller onboarding/storefront reads.",
        )),
        ("Key Risks & Mitigations", (
            "Naming drift: standardize on public.wall_feed_view and update callers.",
            "Vault inserts: align schema or reduce insert payload.",
            "Search null images: add image_best with COALESCE fallback.",
            "RLS/grants: explicit SELECT grants; smoke tests for policies.",
        )),
        ("KPIs For Beta", (
            "TTI to first listing: < 5 minutes on clean device.",
            "Search column error rate: 0% (last 500 searches).",
            "Feed refresh latency: < 2s perceived.",
            "Vault add success: > 99% (last 200 inserts).",
        )),
        ("Tech Highlights", (
            "Supabase + RPCs + materialized views for feed/search.",
            "Flutter ThunderShell app with guarded routing and scanner dev tools.",
            "RLS-first multi-tenant design.",
            "Scripts for migration repair/pull/push and local dev alignment.",
        )),
        ("Run Checklist (Device)", (
            "supabase start (or connect hosted)",
            "flutter clean && flutter pub get",
            "flutter run -d <device-id>",
            "Verify: Search loads (no column errors); Wall feed after refresh/listing; Vault adds without column mismatches.",
        )),
    ]),
}


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(doc: Document, text: str):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)


def main():
    out_dir = Path('docs')
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / 'Grookai_Vault_Investor_Brief.docx'

    doc = Document()
    add_title(doc, CONTENT["title"]) 

    for heading, bullets in CONTENT["sections"]:
        add_heading(doc, heading)
        add_bullets(doc, bullets)

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == '__main__':
    main()

